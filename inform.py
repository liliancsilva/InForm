import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import os
from docx import Document
from copy import deepcopy
import pandas as pd
from datetime import datetime

# ─────────────────────────────────────────────
# LÓGICA DE PREENCHIMENTO
# ─────────────────────────────────────────────

def substituir_texto(paragrafo, substituicoes):
    texto_total = "".join(run.text for run in paragrafo.runs)
    alterado = False
    for chave, valor in substituicoes.items():
        if chave in texto_total:
            texto_total = texto_total.replace(chave, str(valor))
            alterado = True
    if alterado:
        for run in paragrafo.runs:
            run.text = ""
        if paragrafo.runs:
            paragrafo.runs[0].text = texto_total


def substituir_em_celula(celula, substituicoes):
    for p in celula.paragraphs:
        substituir_texto(p, substituicoes)
    for chave, valor in substituicoes.items():
        texto_celula = "\n".join("".join(r.text for r in p.runs) for p in celula.paragraphs)
        if chave not in texto_celula:
            continue
        texto_completo = texto_celula.replace(chave, str(valor))
        linhas = texto_completo.split("\n")
        for i, p in enumerate(celula.paragraphs):
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = linhas[i] if i < len(linhas) else ""


def substituir_em_doc(doc, substituicoes):
    for p in doc.paragraphs:
        substituir_texto(p, substituicoes)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_em_celula(celula, substituicoes)


def encontrar_linha_placeholder(tabela, placeholder):
    for i, linha in enumerate(tabela.rows):
        for celula in linha.cells:
            for p in celula.paragraphs:
                if placeholder in "".join(r.text for r in p.runs):
                    return i
    return None


def duplicar_linha(tabela, idx_linha):
    linha_orig = tabela.rows[idx_linha]._tr
    nova_linha = deepcopy(linha_orig)
    linha_orig.addnext(nova_linha)
    return tabela.rows[idx_linha + 1]


def preencher_tabela_equipe(doc, membros):
    for tabela in doc.tables:
        idx = encontrar_linha_placeholder(tabela, "{{EMAIL}}")
        if idx is None:
            idx = encontrar_linha_placeholder(tabela, "{{NOME}}")
        if idx is None:
            continue
        for _ in range(len(membros) - 1):
            duplicar_linha(tabela, idx)
        for i, membro in enumerate(membros):
            linha = tabela.rows[idx + i]
            subs = {
                "{{EMAIL}}":   membro.get("email", ""),
                "{{NOME}}":    membro.get("nome", ""),
                "{{RG}}":      membro.get("rg", ""),
                "{{CPF}}":     membro.get("cpf", ""),
                "{{EMPRESA}}": membro.get("empresa", ""),
            }
            for celula in linha.cells:
                for p in celula.paragraphs:
                    substituir_texto(p, subs)
        return


def carregar_excel(caminho):
    xls = pd.ExcelFile(caminho)

    df_eq = xls.parse("EQUIPE", header=None)
    membros = []
    for i, row in df_eq.iterrows():
        email   = str(row.iloc[1]).strip() if pd.notna(row.iloc[1]) else ""
        nome    = str(row.iloc[2]).strip() if pd.notna(row.iloc[2]) else ""
        rg      = str(row.iloc[3]).strip() if pd.notna(row.iloc[3]) else ""
        cpf     = str(row.iloc[4]).strip() if pd.notna(row.iloc[4]) else ""
        empresa = str(row.iloc[5]).strip() if pd.notna(row.iloc[5]) else ""
        if not email or email == "nan" or "{{" in email or "@" not in email:
            continue
        membros.append({"email": email, "nome": nome, "rg": rg, "cpf": cpf, "empresa": empresa})

    df_d = xls.parse("DADOS", header=None)
    periodo   = str(df_d.iloc[2, 0]).strip() if pd.notna(df_d.iloc[2, 0]) else ""
    atividade = str(df_d.iloc[4, 0]).strip() if pd.notna(df_d.iloc[4, 0]) else ""

    sites = []
    for i in range(8, len(df_d)):
        site  = str(df_d.iloc[i, 1]).strip() if pd.notna(df_d.iloc[i, 1]) else ""
        infra = str(df_d.iloc[i, 2]).strip() if pd.notna(df_d.iloc[i, 2]) else ""
        end   = str(df_d.iloc[i, 3]).strip() if pd.notna(df_d.iloc[i, 3]) else ""
        if not site or site == "nan" or "{{" in site or "←" in site or "Substitua" in site:
            continue
        infra_norm = infra.strip().lower()
        if infra_norm in ("greenfield", "gf"):
            infra_norm = "GF"
        elif infra_norm in ("rooftop", "rt"):
            infra_norm = "RT"
        sites.append({"site": site, "infra": infra_norm, "endereco": end})

    return [{"nome": "Documento", "sites": sites, "periodo": periodo, "atividade": atividade, "membros": membros}]


def gerar_documentos(arquivo_excel, pasta_pai, log_fn):
    grupos = carregar_excel(arquivo_excel)
    pasta_modelos = os.path.dirname(arquivo_excel)
    pasta_saida   = os.path.join(pasta_pai, "Arquivos Gerados")
    os.makedirs(pasta_saida, exist_ok=True)

    total   = sum(len(g["sites"]) for g in grupos)
    gerados = 0
    erros   = []

    meses = ["janeiro","fevereiro","março","abril","maio","junho",
             "julho","agosto","setembro","outubro","novembro","dezembro"]

    log_fn(f"{'─'*46}")
    log_fn(f"  Total : {total} arquivo(s) a gerar")
    log_fn(f"  Saída : .../{os.path.basename(pasta_pai)}/Arquivos Gerados")
    log_fn(f"{'─'*46}")

    for grupo in grupos:
        for item in grupo["sites"]:
            site     = item["site"]
            infra    = item["infra"]
            endereco = item["endereco"]
            try:
                if infra == "RT":
                    caminho_modelo = os.path.join(pasta_modelos, "Modelo_RT.docx")
                    prefixo = "Modelo - RT"
                else:
                    caminho_modelo = os.path.join(pasta_modelos, "Modelo_GF.docx")
                    prefixo = "Modelo - GF"

                if not os.path.exists(caminho_modelo):
                    raise FileNotFoundError(f"Modelo não encontrado: {os.path.basename(caminho_modelo)}")

                doc  = Document(caminho_modelo)
                hoje = datetime.today()
                data_atual = f"São Paulo, {hoje.day} de {meses[hoje.month-1]} de {hoje.year}."

                subs_doc = {
                    "{{SITE}}":      site,
                    "{{PERIODO}}":   grupo["periodo"],
                    "{{ATIVIDADE}}": grupo["atividade"],
                    "{{ENDERECO}}":  endereco,
                    "{{DATA}}":      data_atual,
                }
                substituir_em_doc(doc, subs_doc)

                if grupo["membros"]:
                    preencher_tabela_equipe(doc, grupo["membros"])

                nome_arquivo = f"{prefixo} {site}.docx"
                doc.save(os.path.join(pasta_saida, nome_arquivo))
                gerados += 1
                log_fn(f"  ✓  {nome_arquivo}")

            except Exception as e:
                erros.append((site, str(e)))
                log_fn(f"  ✗  ERRO em {site}: {e}")

    log_fn(f"{'─'*46}")
    log_fn(f"  Concluído: {gerados}/{total} arquivo(s) gerado(s)")
    if erros:
        log_fn(f"  Erros   : {len(erros)}")
        for site, msg in erros:
            log_fn(f"    - {site}: {msg}")
    log_fn(f"{'─'*46}")
    return gerados, erros


# ─────────────────────────────────────────────
# INTERFACE GRÁFICA
# ─────────────────────────────────────────────

COR_FUNDO  = "#F5F5F5"
COR_PAINEL = "#FFFFFF"
COR_BORDA  = "#DDDDDD"
COR_AZUL   = "#1F4E79"
COR_AZUL_H = "#2E75B6"
COR_TEXTO  = "#222222"
COR_CINZA  = "#666666"
COR_LOG_BG = "#1E1E1E"
COR_LOG_FG = "#D4D4D4"
FONTE      = "Segoe UI"


class InFormApp:
    def __init__(self, root):
        self.root = root
        self.root.title("InForm — Preenchedor de Formulários de Acesso")
        self.root.geometry("680x520")
        self.root.resizable(False, False)
        self.root.configure(bg=COR_FUNDO)

        self.excel_path = tk.StringVar()
        self.saida_path = tk.StringVar()
        self.rodando    = False

        self._build_ui()

    def _build_ui(self):
        # ── Cabeçalho ──
        header = tk.Frame(self.root, bg=COR_AZUL, height=60)
        header.pack(fill="x")
        header.pack_propagate(False)

        tk.Label(header, text="InForm", font=(FONTE, 18, "bold"),
                 bg=COR_AZUL, fg="white").pack(side="left", padx=20, pady=10)
        tk.Label(header, text="Preenchedor de Formulários de Acesso",
                 font=(FONTE, 10), bg=COR_AZUL, fg="#BDD7EE").pack(side="left", pady=16)

        # ── Corpo ──
        corpo = tk.Frame(self.root, bg=COR_FUNDO, padx=24, pady=16)
        corpo.pack(fill="both", expand=True)
        corpo.columnconfigure(1, weight=1)

        # Arquivo Excel
        self._campo(corpo, "Arquivo de dados (.xlsx)", self.excel_path,
                    self._selecionar_excel, "Selecione o arquivo .xlsx...", row=0)

        # Pasta de saída
        self._campo(corpo, "Pasta de saída (será criada subpasta 'Arquivos Gerados')",
                    self.saida_path, self._selecionar_saida,
                    "Selecione onde salvar os arquivos gerados...", row=2)

        # Botão gerar
        self.btn_gerar = tk.Button(
            corpo,
            text="▶  Gerar Documentos",
            font=(FONTE, 11, "bold"),
            bg=COR_AZUL, fg="white",
            activebackground=COR_AZUL_H, activeforeground="white",
            relief="flat", cursor="hand2",
            padx=20, pady=10,
            command=self._iniciar_geracao
        )
        self.btn_gerar.grid(row=4, column=0, columnspan=3, pady=(16, 8), sticky="ew")

        # Log
        tk.Label(corpo, text="Log de execução", font=(FONTE, 9),
                 bg=COR_FUNDO, fg=COR_CINZA).grid(row=5, column=0, columnspan=3, sticky="w")

        log_frame = tk.Frame(corpo, bg=COR_LOG_BG, bd=1, relief="solid")
        log_frame.grid(row=6, column=0, columnspan=3, sticky="nsew")
        corpo.rowconfigure(6, weight=1)

        self.log_text = tk.Text(
            log_frame,
            font=("Consolas", 9),
            bg=COR_LOG_BG, fg=COR_LOG_FG,
            relief="flat", bd=6,
            state="disabled",
            wrap="word",
            height=12
        )
        self.log_text.pack(side="left", fill="both", expand=True)

        scroll = tk.Scrollbar(log_frame, command=self.log_text.yview)
        scroll.pack(side="right", fill="y")
        self.log_text.configure(yscrollcommand=scroll.set)

        self.log_text.tag_config("ok",   foreground="#6BCB77")
        self.log_text.tag_config("erro", foreground="#FF6B6B")
        self.log_text.tag_config("info", foreground="#BDD7EE")
        self.log_text.tag_config("sep",  foreground="#555555")

        # Status bar
        self.status_var = tk.StringVar(value="Pronto.")
        tk.Label(self.root, textvariable=self.status_var,
                 font=(FONTE, 8), bg=COR_BORDA, fg=COR_CINZA,
                 anchor="w", padx=10).pack(fill="x", side="bottom")

    def _campo(self, parent, label, var, cmd, placeholder, row):
        tk.Label(parent, text=label, font=(FONTE, 9, "bold"),
                 bg=COR_FUNDO, fg=COR_TEXTO).grid(row=row, column=0, columnspan=3,
                                                    sticky="w", pady=(8, 2))
        tk.Entry(parent, textvariable=var, font=(FONTE, 9),
                 bg=COR_PAINEL, fg=COR_TEXTO, relief="solid", bd=1,
                 highlightthickness=1, highlightbackground=COR_BORDA,
                 highlightcolor=COR_AZUL_H
                 ).grid(row=row+1, column=0, columnspan=2, sticky="ew", ipady=6, padx=(0, 6))
        tk.Button(parent, text="Procurar", font=(FONTE, 9),
                  bg=COR_AZUL, fg="white",
                  activebackground=COR_AZUL_H, activeforeground="white",
                  relief="flat", cursor="hand2", padx=10,
                  command=cmd
                  ).grid(row=row+1, column=2, sticky="ew", ipady=6)

    def _selecionar_excel(self):
        path = filedialog.askopenfilename(
            title="Selecionar arquivo Excel",
            filetypes=[("Excel", "*.xlsx *.xls"), ("Todos", "*.*")]
        )
        if path:
            self.excel_path.set(path)

    def _selecionar_saida(self):
        path = filedialog.askdirectory(title="Selecionar pasta de saída")
        if path:
            self.saida_path.set(path)

    def _log(self, msg):
        self.log_text.configure(state="normal")
        if msg.startswith("  ✓"):
            tag = "ok"
        elif msg.startswith("  ✗") or "ERRO" in msg:
            tag = "erro"
        elif "─" in msg:
            tag = "sep"
        else:
            tag = "info"
        self.log_text.insert("end", msg + "\n", tag)
        self.log_text.see("end")
        self.log_text.configure(state="disabled")

    def _limpar_log(self):
        self.log_text.configure(state="normal")
        self.log_text.delete("1.0", "end")
        self.log_text.configure(state="disabled")

    def _iniciar_geracao(self):
        if self.rodando:
            return

        excel = self.excel_path.get().strip()
        saida = self.saida_path.get().strip()

        if not excel:
            messagebox.showwarning("Atenção", "Selecione o arquivo Excel de dados.")
            return
        if not os.path.exists(excel):
            messagebox.showerror("Erro", f"Arquivo não encontrado:\n{excel}")
            return
        if not saida:
            messagebox.showwarning("Atenção", "Selecione a pasta de saída.")
            return

        self._limpar_log()
        self.rodando = True
        self.btn_gerar.configure(state="disabled", text="⏳  Gerando...")
        self.status_var.set("Gerando documentos...")

        thread = threading.Thread(target=self._executar, args=(excel, saida), daemon=True)
        thread.start()

    def _executar(self, excel, saida):
        try:
            gerados, erros = gerar_documentos(excel, saida, self._log)
            pasta_final = os.path.join(saida, "Arquivos Gerados")

            if erros:
                self.status_var.set(f"Concluído com {len(erros)} erro(s). {gerados} arquivo(s) gerado(s).")
            else:
                self.status_var.set(f"✓ {gerados} arquivo(s) gerado(s) com sucesso!")
                self.root.after(100, lambda: messagebox.showinfo(
                    "Concluído", f"{gerados} documento(s) gerado(s) com sucesso!\n\nSalvos em:\n{pasta_final}"
                ))

        except Exception as e:
            self._log(f"  ✗  Erro inesperado: {e}")
            self.status_var.set("Erro durante a execução.")
            self.root.after(100, lambda: messagebox.showerror("Erro", str(e)))

        finally:
            self.rodando = False
            self.root.after(0, lambda: self.btn_gerar.configure(
                state="normal", text="▶  Gerar Documentos"
            ))


# ─────────────────────────────────────────────
# EXECUÇÃO
# ─────────────────────────────────────────────

if __name__ == "__main__":
    root = tk.Tk()
    app  = InFormApp(root)
    root.mainloop()
